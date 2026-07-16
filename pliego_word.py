"""Exportación del pliego de condiciones a Word (.docx) — versión profesional.

Genera un documento con portada, índice automático (campo TOC nativo de
Word), cada sección con estilos Heading 1/Heading 2/Normal, numeración de
páginas y pie de página con el nombre del proyecto.

Convierte el HTML simple que genera la IA (p, ul/li, table, strong/em, y
opcionalmente h3 para subapartados) en párrafos, listas y tablas nativas de
python-docx.
"""

from datetime import datetime
from html.parser import HTMLParser
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


class _HTMLToDocxBlocks(HTMLParser):
    """Primer paso: convierte el HTML de una sección en una lista de bloques
    simples (encabezado, párrafo, ítem de lista, tabla) que luego se
    renderizan en el documento."""

    def __init__(self):
        super().__init__()
        self.blocks = []
        self._runs = None
        self._bold = False
        self._italic = False
        self._list_ordered = False
        self._table_rows = None
        self._current_row = None
        self._current_cell = None
        self._heading_level = None

    def handle_starttag(self, tag, attrs):
        if tag == "p":
            self._runs = []
        elif tag == "li":
            self._runs = []
        elif tag == "ol":
            self._list_ordered = True
        elif tag == "ul":
            self._list_ordered = False
        elif tag in ("h2", "h3", "h4"):
            self._runs = []
            self._heading_level = {"h2": 2, "h3": 2, "h4": 3}[tag]
        elif tag == "table":
            self._table_rows = []
        elif tag == "tr":
            self._current_row = []
        elif tag in ("td", "th"):
            self._current_cell = []
        elif tag in ("strong", "b"):
            self._bold = True
        elif tag in ("em", "i"):
            self._italic = True

    def handle_endtag(self, tag):
        if tag == "p":
            if self._runs:
                self.blocks.append({"type": "p", "runs": self._runs})
            self._runs = None
        elif tag == "li":
            if self._runs:
                self.blocks.append({"type": "li", "runs": self._runs, "ordered": self._list_ordered})
            self._runs = None
        elif tag in ("h2", "h3", "h4"):
            if self._runs:
                self.blocks.append({"type": "heading", "level": self._heading_level, "runs": self._runs})
            self._runs = None
            self._heading_level = None
        elif tag == "table":
            if self._table_rows:
                self.blocks.append({"type": "table", "rows": self._table_rows})
            self._table_rows = None
        elif tag == "tr":
            if self._current_row is not None and self._table_rows is not None:
                self._table_rows.append(self._current_row)
            self._current_row = None
        elif tag in ("td", "th"):
            if self._current_row is not None:
                self._current_row.append("".join(self._current_cell) if self._current_cell else "")
            self._current_cell = None
        elif tag in ("strong", "b"):
            self._bold = False
        elif tag in ("em", "i"):
            self._italic = False

    def handle_data(self, data):
        if self._current_cell is not None:
            self._current_cell.append(data)
        elif self._runs is not None:
            self._runs.append((data, self._bold, self._italic))


def _render_blocks(document: Document, blocks: list):
    for block in blocks:
        if block["type"] == "heading":
            document.add_heading(
                "".join(text for text, _, _ in block["runs"]),
                level=block["level"],
            )
        elif block["type"] in ("p", "li"):
            style = None
            if block["type"] == "li":
                style = "List Number" if block.get("ordered") else "List Bullet"
            paragraph = document.add_paragraph(style=style)
            for text, bold, italic in block["runs"]:
                run = paragraph.add_run(text)
                run.bold = bold
                run.italic = italic
        elif block["type"] == "table":
            rows = block["rows"]
            if not rows:
                continue
            n_cols = max(len(r) for r in rows)
            table = document.add_table(rows=len(rows), cols=n_cols)
            table.style = "Light Grid Accent 1"
            for i, row in enumerate(rows):
                for j in range(n_cols):
                    table.cell(i, j).text = row[j] if j < len(row) else ""


def _add_field(paragraph, instr_text: str):
    """Inserta un campo de Word (ej. PAGE, NUMPAGES) que se calcula al abrir el documento."""
    run = paragraph.add_run()
    r = run._r

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instr_text

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    r.append(begin)
    r.append(instr)
    r.append(end)


def _add_toc(document: Document):
    """Inserta un índice automático (campo TOC nativo de Word).

    Word no calcula el contenido al generarlo por código: al abrir el
    documento, Word pedirá actualizar los campos (o lo hará solo, gracias a
    `_force_update_fields_on_open`), y entonces rellenará el índice real con
    los títulos (Heading 1/2) y su número de página.
    """
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    r = run._r

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "Haz clic derecho sobre el índice y selecciona «Actualizar campos» para generarlo."

    placeholder_run = OxmlElement("w:r")
    placeholder_run.append(placeholder)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    r.append(begin)
    r.append(instr)

    p = paragraph._p
    p.append(separate)
    p.append(placeholder_run)
    p.append(end)


def _force_update_fields_on_open(document: Document):
    """Marca el documento para que Word recalcule los campos (índice,
    números de página) automáticamente en cuanto se abre."""
    settings = document.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


def _add_portada(document: Document, datos: dict):
    for _ in range(4):
        document.add_paragraph()

    titulo = document.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("PLIEGO DE CONDICIONES TÉCNICAS Y ADMINISTRATIVAS")
    run.bold = True
    run.font.size = Pt(22)

    subtitulo = document.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitulo.add_run(datos.get("nombre_proyecto", ""))
    run_sub.italic = True
    run_sub.font.size = Pt(15)

    document.add_paragraph()
    document.add_paragraph()

    campos = [
        ("Arquitecto", datos.get("arquitecto", "")),
        ("Municipio", f"{datos.get('municipio', '')} ({datos.get('provincia', '')})"),
        ("Tipo de obra", datos.get("tipo_obra", "")),
        ("Calidad de acabados", datos.get("calidad_acabados", "")),
        ("Fecha de generación", datetime.now().strftime("%d/%m/%Y")),
    ]

    tabla = document.add_table(rows=0, cols=2)
    tabla.autofit = True
    for etiqueta, valor in campos:
        fila = tabla.add_row()
        celda_etiqueta = fila.cells[0]
        celda_etiqueta.text = etiqueta
        celda_etiqueta.paragraphs[0].runs[0].bold = True
        fila.cells[1].text = str(valor)


def _add_pie_pagina(document: Document, nombre_proyecto: str):
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(f"{nombre_proyecto}   ·   Página ")
    _add_field(paragraph, "PAGE")
    paragraph.add_run(" de ")
    _add_field(paragraph, "NUMPAGES")


def construir_word(datos: dict, secciones: list) -> bytes:
    """Genera el Word profesional del pliego: portada, índice automático,
    secciones con Heading 1/2/Normal, pie de página y numeración."""
    document = Document()

    _add_portada(document, datos)
    document.add_page_break()

    document.add_heading("Índice", level=1)
    _add_toc(document)
    document.add_page_break()

    for sec in secciones:
        document.add_heading(f"{sec['codigo']}. {sec['titulo']}", level=1)
        parser = _HTMLToDocxBlocks()
        parser.feed(sec.get("contenido_html", ""))
        _render_blocks(document, parser.blocks)

    _add_pie_pagina(document, datos.get("nombre_proyecto", ""))
    _force_update_fields_on_open(document)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
